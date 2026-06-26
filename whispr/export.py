"""Export transcripts and translations to Microsoft Word (``.docx``).

A small wrapper over ``python-docx`` (imported lazily, like the rest of the
package's optional backends) so an operator can hand off a clean Word document
instead of pasting plain text. Diarized transcripts keep their speaker labels in
bold so turns are easy to scan.
"""

from __future__ import annotations

from pathlib import Path
from typing import Dict, Optional, Union

from .transcription import TranscriptionResult

PathLike = Union[str, Path]


def _document():
    """Return a new python-docx Document, with a clear error if it's missing."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is not installed. Install the GUI extra with:  "
            "pip install 'silvance-whisper[gui]'"
        ) from exc
    return Document()


def text_to_docx(text: str, path: PathLike, *, title: Optional[str] = None) -> Path:
    """Write ``text`` to a ``.docx`` at ``path``, one paragraph per blank-line block.

    Used for translation output. Blank lines separate paragraphs; single newlines
    inside a block become line breaks within a paragraph.
    """
    document = _document()
    if title:
        document.add_heading(title, level=1)
    # Split on blank lines into paragraphs; keep intra-paragraph line breaks.
    blocks = text.replace("\r\n", "\n").split("\n\n")
    for block in blocks:
        paragraph = document.add_paragraph()
        lines = block.split("\n")
        for line_index, line in enumerate(lines):
            if line_index:
                paragraph.add_run().add_break()
            paragraph.add_run(line)
    out = Path(path)
    document.save(str(out))
    return out


def transcript_to_docx(
    result: TranscriptionResult,
    path: PathLike,
    speaker_names: Optional[Dict[str, str]] = None,
    *,
    blank_lines: bool = True,
    title: Optional[str] = None,
) -> Path:
    """Write a transcript to a ``.docx``.

    When diarized, each paragraph starts with the speaker label in bold
    (optionally remapped via ``speaker_names``); otherwise one paragraph per
    segment. ``blank_lines`` is accepted for signature parity with the text
    exporters - Word paragraphs already render with spacing between them.
    """
    document = _document()
    if title:
        document.add_heading(title, level=1)

    if not result.has_speakers:
        segments = result.segments
        if segments:
            for segment in segments:
                document.add_paragraph(segment.text)
        else:
            for line in result.text.split("\n"):
                document.add_paragraph(line)
    else:
        for segment in result.segments:
            label = result._speaker_label(segment.speaker, speaker_names)
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"[{label}] ")
            run.bold = True
            paragraph.add_run(segment.text)

    out = Path(path)
    document.save(str(out))
    return out
