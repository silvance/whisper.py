import importlib.util

import pytest

from whispr.export import text_to_docx, transcript_to_docx
from whispr.transcription import Segment, TranscriptionResult

_HAS_DOCX = importlib.util.find_spec("docx") is not None


@pytest.mark.skipif(_HAS_DOCX, reason="python-docx installed; error path not exercised")
def test_text_to_docx_requires_python_docx(tmp_path):
    with pytest.raises(RuntimeError, match="python-docx is not installed"):
        text_to_docx("hello", tmp_path / "out.docx")


@pytest.mark.skipif(_HAS_DOCX, reason="python-docx installed; error path not exercised")
def test_transcript_to_docx_requires_python_docx(tmp_path):
    result = TranscriptionResult(
        text="hi", language="en", language_probability=1.0, duration=1.0
    )
    with pytest.raises(RuntimeError, match="python-docx is not installed"):
        transcript_to_docx(result, tmp_path / "out.docx")


@pytest.mark.skipif(not _HAS_DOCX, reason="python-docx not installed")
def test_text_to_docx_writes_paragraphs(tmp_path):
    from docx import Document

    out = text_to_docx("first block\nsecond line\n\nthird block", tmp_path / "t.docx")
    assert out.exists()
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "first block" in texts[0]
    assert "third block" in "\n".join(texts)


@pytest.mark.skipif(not _HAS_DOCX, reason="python-docx not installed")
def test_transcript_to_docx_with_speakers(tmp_path):
    from docx import Document

    result = TranscriptionResult(
        text="hello\nworld",
        language="en",
        language_probability=1.0,
        duration=2.0,
        segments=[
            Segment(start=0.0, end=1.0, text="hello", speaker="SPEAKER_00"),
            Segment(start=1.0, end=2.0, text="world", speaker="SPEAKER_01"),
        ],
    )
    out = transcript_to_docx(result, tmp_path / "t.docx", {"SPEAKER_00": "Alice"})
    doc = Document(str(out))
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "[Alice] hello" in joined
    assert "[SPEAKER_01] world" in joined
